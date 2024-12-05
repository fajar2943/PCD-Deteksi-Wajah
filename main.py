import cv2
import os, shutil
from PIL import Image
from docx import Document
from docx.shared import Inches

# Path ke folder dataset dan file cascade XML
dataset_folder = "train"  # Sesuaikan dengan lokasi folder dataset Anda
output_folder = "output"  # Folder untuk menyimpan hasil
cascade_path = "face_ref.xml"  # Path file haarcascade
output_docx = "output_faces.docx"  # Nama file Word output

# Load classifier wajah
face_cascade = cv2.CascadeClassifier(cascade_path)

# pengolahan citra dengan metode grayscale
def gray_detect_face(image_path, output_path):
    # Membaca gambar
    image = cv2.imread(image_path)
    if image is None:
        print(f"Error: Tidak dapat memuat gambar {image_path}")
        return
    
    # Mengubah ke dalam skala abu-abu (grayscale)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    save_box_image(image, gray, output_path)

# pengolahan citra dengan metode gaussian blur
def gaussian_detect_face(image_path, output_path):
    # Membaca gambar
    image = cv2.imread(image_path)
    if image is None:
        print(f"Error: Tidak dapat memuat gambar {image_path}")
        return
    
    # memperbaiki citra menggunakan gaussian blur dengan menghilangkan noise
    gaussian = cv2.GaussianBlur(image, (5, 5), 0)

    save_box_image(image, gaussian, output_path)
    
def save_box_image(image, image_processing, output_path):    
    # Deteksi wajah
    faces = face_cascade.detectMultiScale(image_processing, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))
    # Menggambar kotak merah di setiap wajah yang terdeteksi
    for (x, y, w, h) in faces:
        cv2.rectangle(image_processing, (x, y), (x + w, y + h), (0, 0, 255), 2)

    # Menyimpan hasil gambar dalam format .jpg
    cv2.imwrite(output_path, image_processing)
    print(f"Hasil disimpan di: {output_path}")

def main():    
    # Cek apakah file XML ada
    if face_cascade.empty():
        print("Error: Gagal memuat haarcascade.")
        exit()
    try:
        # Membuat folder output jika belum ada
        if os.path.exists(output_folder):
            shutil.rmtree(output_folder) #hapus folder output
        os.makedirs(output_folder) #buat folder output
        os.makedirs(output_folder+'/gray') #buat folder output
        os.makedirs(output_folder+'/gaussian') #buat folder output
        
    except:
        print("Coba tutup word nya dulu!!!")
        exit()
    # List untuk menyimpan path gambar hasil deteksi
    output_images = []
    # Iterasi setiap gambar di folder dataset
    for s_folders_name in os.listdir(dataset_folder): #s_folders_name adalah nama dari folder didalam folder train
        s_folder = os.path.join(dataset_folder, s_folders_name)
        for filename in os.listdir(s_folder): #buka folder yang ada didalam folder train
            file_path = os.path.join(s_folder, filename)
            # print(file_path)
            # exit()
            # Cek ekstensi file untuk memastikan hanya file gambar yang diproses
            if os.path.isfile(file_path) and filename.lower().endswith(('.jpg', '.jpeg', '.png', '.ppm')):
                gray_output_path = os.path.join(output_folder+'/gray', os.path.splitext(filename)[0] + ".jpg")  # Ganti ekstensi dengan .jpg
                print(f"Mendeteksi wajah pada {filename}")
                gray_detect_face(file_path, gray_output_path)

                gaussian_output_path = os.path.join(output_folder+'/gaussian', os.path.splitext(filename)[0] + ".jpg")  # Ganti ekstensi dengan .jpg
                print(f"Mendeteksi wajah pada {filename}")
                gaussian_detect_face(file_path, gaussian_output_path)

                output_images.append((gray_output_path, filename))  # Simpan tuple (path gambar, nama file)
                output_images.append((gaussian_output_path, filename))  # Simpan tuple (path gambar, nama file)


    # print(output_images)
    # exit()

    #Menyimpan hasil deteksi wajah ke satu file Word
    if output_images:
        doc = Document()
        title = doc.add_heading("Hasil Deteksi Wajah, Matkul PCD", level=1)
        titleStyle = title.runs[0]
        titleStyle.font.name = 'Times New Roman'
        titleStyle.bold = True

        nama = doc.add_heading("Fajar Lintang Gumilang (A11.2022.14832)", level=5)
        namaStyle = nama.runs[0]
        namaStyle.font.name = 'Times New Roman'
        namaStyle.bold = True

        doc.add_paragraph("", style="Caption")
        doc.add_paragraph("", style="Caption")

        print("Membuat file Docx...")

        i = 0
        for img_path, filename in output_images:
            # Menambahkan gambar ke dokumen
            doc.add_picture(img_path, width=Inches(4))  # Sesuaikan ukuran lebar gambar
            # Menambahkan keterangan nama file
            doc.add_paragraph(f"Nama file: {filename}", style="Caption")
            doc.add_paragraph(f"Path file: {img_path}", style="Caption")
            doc.add_paragraph(f"Hasil analisa: ", style="Caption")
            i += 1

        # Menyimpan dokumen Word
        doc.save(os.path.join(output_folder, output_docx))
        print(f"Hasil semua gambar disimpan dalam file Word: {output_docx}")
    else:
        print("Tidak ada gambar yang ditemukan untuk disimpan dalam file Word.")



if __name__ == '__main__':
    main()
